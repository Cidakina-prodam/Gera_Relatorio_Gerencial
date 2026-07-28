"""
Conferência GDS-1 — Horas, Receita, Colaboradores e Despesa
=============================================================
App de apoio para conferir rapidamente os dados do Portal GDS-1 sem precisar
montar o relatório completo em Word. Aplica a correção de horas por núcleo
(prioriza o detalhamento por núcleo sobre o campo-resumo do contrato).

Como rodar localmente:
    pip install -r requirements.txt
    streamlit run app.py

Como publicar (gratuito):
    1. Suba esta pasta para um repositório no GitHub.
    2. Acesse https://share.streamlit.io, conecte sua conta GitHub.
    3. Escolha o repositório e o arquivo app.py. Deploy.
"""

import io
import json
from collections import Counter, defaultdict

import pandas as pd
import streamlit as st
from docx import Document

st.set_page_config(page_title="Conferência GDS-1", page_icon="📊", layout="wide")

# ----------------------------------------------------------------------
# Helpers de cálculo (mesma lógica corrigida usada no relatório em Word)
# ----------------------------------------------------------------------

def num(v):
    if v in ("", None):
        return 0.0
    return float(v)


def horas_contrato_mes(lancamentos_mes: dict, contrato_id: str):
    """Prioriza o detalhamento por núcleo ('{id}_{NUCLEO}') sobre o campo-resumo ('{id}')."""
    prefixo = f"{contrato_id}_"
    splits = {k: v for k, v in lancamentos_mes.items() if k.startswith(prefixo)}
    if splits:
        posH = sum(num(v.get("posH")) for v in splits.values())
        gfpH = sum(num(v.get("gfpH")) for v in splits.values())
    else:
        bare = lancamentos_mes.get(contrato_id, {})
        posH = num(bare.get("posH"))
        gfpH = num(bare.get("gfpH"))
    return posH, gfpH


def nucleo_do_split(chave: str) -> str:
    return chave.split("_", 1)[1]


def normalizar(txt: str) -> str:
    return "".join(ch for ch in (txt or "").upper() if ch.isalnum())


def agrupar_familias(contratos: dict) -> dict:
    """
    Agrupa contratos duplicados (mesmo contrato real recriado com novo id no Portal)
    usando (secretaria + número do contrato normalizado) como chave.
    Retorna {chave_normalizada: {'nome': str, 'ids': [...]}}
    """
    familias = {}
    for cid, c in contratos.items():
        chave = normalizar(c.get("secretaria", "")) + "|" + normalizar(c.get("contrato", ""))
        if chave not in familias:
            familias[chave] = {"nome": f"{c.get('secretaria')} — {c.get('contrato')}", "ids": []}
        familias[chave]["ids"].append(cid)
    return familias


def processar_despesa(backup: dict, meses: list) -> pd.DataFrame:
    """Direto = colaboradores tipo DIRETO; Indireto = equipeIndireta (não soma tipo INDIRETO
    dentro de custos, pra não contar a mesma pessoa duas vezes); Terceiros = demais tipos."""
    linhas = []
    for mes in meses:
        bloco = backup.get("lancamentos", {}).get(mes, {})
        custos = bloco.get("custos", {})
        equipe_ind = bloco.get("equipeIndireta", [])
        direto = terceiros = 0.0
        for nuc, v in custos.items():
            for c in v.get("colaboradores", []):
                if c.get("tipo") == "DIRETO":
                    direto += num(c.get("custo"))
                elif c.get("tipo") == "INDIRETO":
                    continue
                else:
                    terceiros += num(c.get("custo"))
        indireto = sum(num(c.get("custo")) for c in equipe_ind)
        linhas.append(
            dict(mes=mes, direto=direto, indireto=indireto, terceiros=terceiros, total=direto + indireto + terceiros)
        )
    return pd.DataFrame(linhas).set_index("mes")


MESES_RELATORIO = ["2026-01", "2026-02", "2026-03", "2026-04", "2026-05", "2026-06"]


@st.cache_data(show_spinner=False)
def processar_portal(backup_bytes: bytes):
    backup = json.loads(backup_bytes)
    contratos = {str(c["id"]): c for c in backup["contratos"]}
    meses = [m for m in MESES_RELATORIO if m in backup["lancamentos"]]

    linhas = []  # detalhe por contrato/mes
    for mes in meses:
        lan = backup["lancamentos"][mes].get("contratos", {})
        for cid, c in contratos.items():
            posH, gfpH = horas_contrato_mes(lan, cid)
            if posH == 0 and gfpH == 0:
                continue
            vlrHora = c.get("vlrHora", 0)
            linhas.append(
                dict(
                    mes=mes,
                    contrato_id=cid,
                    secretaria=c.get("secretaria"),
                    contrato=c.get("contrato"),
                    nucleo=c.get("nucleo"),
                    status=c.get("status"),
                    vlrHora=vlrHora,
                    posH=posH,
                    gfpH=gfpH,
                    posV=posH * vlrHora,
                    gfpV=gfpH * vlrHora,
                )
            )
    df = pd.DataFrame(linhas)

    # núcleo mensal usando o núcleo REAL de cada split (não o núcleo primário do contrato)
    nucleo_rows = []
    for mes in meses:
        lan = backup["lancamentos"][mes].get("contratos", {})
        tmp = defaultdict(float)
        for cid, c in contratos.items():
            splits = {k: v for k, v in lan.items() if k.startswith(cid + "_")}
            bare = lan.get(cid, {})
            if splits:
                for k, v in splits.items():
                    tmp[nucleo_do_split(k)] += num(v.get("posH"))
            else:
                tmp[c["nucleo"]] += num(bare.get("posH"))
        for n, h in tmp.items():
            if h:
                nucleo_rows.append(dict(mes=mes, nucleo=n, posH=h))
    df_nucleo = pd.DataFrame(nucleo_rows)

    return backup, contratos, meses, df, df_nucleo


@st.cache_data(show_spinner=False)
def processar_colaboradores(colab_bytes: bytes):
    d = json.loads(colab_bytes)
    colabs = d["colaboradores"]
    ativos = [c for c in colabs if not c.get("dataSaida")]
    return colabs, ativos


@st.cache_data(show_spinner=False)
def processar_consumo_csv(csv_bytes: bytes):
    df = pd.read_csv(io.BytesIO(csv_bytes), sep=";", encoding="utf-8-sig")
    return df


def fmt_h(v):
    return f"{v:,.1f}".replace(",", "X").replace(".", ",").replace("X", ".")


def fmt_money(v):
    return "R$ " + f"{v:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")


# ----------------------------------------------------------------------
# Interface
# ----------------------------------------------------------------------

st.title("📊 Conferência GDS-1")
st.caption(
    "Sobe os arquivos exportados do Portal GDS-1 (e opcionalmente Colaboradores / OS / Consumo) "
    "e confere na hora as horas, receita, equipe e despesa — já com a correção de horas por núcleo aplicada."
)

with st.sidebar:
    st.header("Arquivos")
    f_portal = st.file_uploader("backup_portal_gds_*.json", type="json", key="portal")
    f_colab = st.file_uploader("backup_colaboradores_*.json (opcional)", type="json", key="colab")
    f_os = st.file_uploader("dados_os_prodam_*.json (opcional)", type="json", key="os")
    f_consumo = st.file_uploader("consumo_os_prodam_*.csv (opcional)", type="csv", key="consumo")
    st.markdown("---")
    st.caption(
        "Correção aplicada: horas são somadas a partir do detalhamento por núcleo "
        "(`contrato_NUCLEO`) sempre que existir, em vez do campo-resumo do contrato "
        "— que pode ficar desatualizado."
    )
    st.caption("📅 Período coberto: Jan-Jun/2026 (mesmo período do template do relatório).")

if not f_portal:
    st.info("⬅️ Suba pelo menos o `backup_portal_gds_*.json` na barra lateral para começar.")
    st.stop()

backup, contratos, meses, df, df_nucleo = processar_portal(f_portal.getvalue())

meses_label = {m: pd.to_datetime(m + "-01").strftime("%b/%y").capitalize() for m in meses}

# ---- KPIs ----
tot_posH = df["posH"].sum()
tot_gfpH = df["gfpH"].sum()
tot_posV = df["posV"].sum()
tot_gfpV = df["gfpV"].sum()

c1, c2, c3, c4 = st.columns(4)
c1.metric("Horas GDS-1 (total)", fmt_h(tot_posH))
c2.metric("Receita GDS-1", fmt_money(tot_posV))
c3.metric("Horas GFP (total)", fmt_h(tot_gfpH))
c4.metric("Receita GFP", fmt_money(tot_gfpV))

st.markdown("---")

# ---- Resumo mensal ----
st.subheader("Resumo mensal — GDS-1 x GFP")
resumo = (
    df.groupby("mes")[["posH", "gfpH", "posV", "gfpV"]]
    .sum()
    .reindex(meses)
    .fillna(0)
)
resumo_fmt = pd.DataFrame(
    {
        "Mês": [meses_label[m] for m in resumo.index],
        "Horas GDS-1": resumo["posH"].map(fmt_h),
        "Receita GDS-1": resumo["posV"].map(fmt_money),
        "Horas GFP": resumo["gfpH"].map(fmt_h),
        "Receita GFP": resumo["gfpV"].map(fmt_money),
    }
)
st.dataframe(resumo_fmt, hide_index=True, use_container_width=True)

# ---- Núcleo mensal ----
st.subheader("Horas GDS-1 por núcleo")
if not df_nucleo.empty:
    piv = df_nucleo.pivot_table(index="nucleo", columns="mes", values="posH", aggfunc="sum").reindex(
        columns=meses
    ).fillna(0)
    piv["Total"] = piv.sum(axis=1)
    piv_fmt = piv.copy()
    for c in piv_fmt.columns:
        piv_fmt[c] = piv_fmt[c].map(fmt_h)
    piv_fmt.columns = [meses_label.get(c, c) for c in piv.columns]
    st.dataframe(piv_fmt, use_container_width=True)
else:
    st.caption("Sem dados de núcleo para o período.")

# ---- Saldo por contrato ----
st.subheader("Saldo de horas por contrato")
familias = agrupar_familias(contratos)
saldo_rows = []
familias_calc = {}  # chave -> dict com posH/gfpH/posV/gfpV por mes + horasTotal
for chave, fam in familias.items():
    ids = fam["ids"]
    acum = df.loc[df["contrato_id"].isin(ids), "posH"].sum()
    # horasTotal do id mais recente (maior valor numérico de id = mais recente/ativo, heurística razoável)
    ids_ordenados = sorted(ids, key=lambda x: int(x) if x.isdigit() else 0)
    ref_id = ids_ordenados[-1]
    horasTotal = contratos[ref_id].get("horasTotal", 0)
    por_mes = {
        m: dict(
            posH=df.loc[(df["contrato_id"].isin(ids)) & (df["mes"] == m), "posH"].sum(),
            gfpH=df.loc[(df["contrato_id"].isin(ids)) & (df["mes"] == m), "gfpH"].sum(),
            posV=df.loc[(df["contrato_id"].isin(ids)) & (df["mes"] == m), "posV"].sum(),
            gfpV=df.loc[(df["contrato_id"].isin(ids)) & (df["mes"] == m), "gfpV"].sum(),
        )
        for m in meses
    }
    familias_calc[chave] = dict(nome=fam["nome"], horasTotal=horasTotal, por_mes=por_mes, acum=acum)
    if horasTotal == 0 and acum == 0:
        continue
    saldo = horasTotal - acum
    pct = acum / horasTotal * 100 if horasTotal else 0
    saldo_rows.append(
        dict(
            Contrato=fam["nome"],
            **{
                "Horas Totais": fmt_h(horasTotal),
                "Acumulado": fmt_h(acum),
                "Saldo": fmt_h(saldo),
                "% Executado": f"{pct:.1f}%".replace(".", ","),
            },
        )
    )
st.dataframe(pd.DataFrame(saldo_rows), hide_index=True, use_container_width=True)
st.caption(
    "Contratos renovados/recriados no Portal (novo id para o mesmo número de contrato) já são "
    "somados automaticamente numa linha só."
)

st.markdown("---")

# ---- Despesa e Margem ----
st.subheader("Receita × Despesa × Margem")
despesa = processar_despesa(backup, meses)
margem_df = pd.DataFrame(index=meses)
margem_df["receita"] = resumo["posV"]
margem_df["despesa"] = despesa["total"]
margem_df["margem"] = margem_df["receita"] - margem_df["despesa"]
margem_df["margem_pct"] = margem_df["margem"] / margem_df["receita"] * 100

margem_gfp_df = pd.DataFrame(index=meses)
margem_gfp_df["receita_gfp"] = resumo["gfpV"]
margem_gfp_df["despesa"] = despesa["total"]
margem_gfp_df["margem"] = margem_gfp_df["receita_gfp"] - margem_gfp_df["despesa"]
margem_gfp_df["margem_pct"] = margem_gfp_df["margem"] / margem_gfp_df["receita_gfp"] * 100

tot_receita = margem_df["receita"].sum()
tot_despesa = margem_df["despesa"].sum()
tot_margem = tot_receita - tot_despesa
tot_receita_gfp = margem_gfp_df["receita_gfp"].sum()
tot_margem_gfp = tot_receita_gfp - tot_despesa

c1, c2, c3 = st.columns(3)
c1.metric("Receita GDS-1", fmt_money(tot_receita))
c2.metric("Despesa Total", fmt_money(tot_despesa))
c3.metric("Margem GDS-1", f"{fmt_money(tot_margem)} ({tot_margem/tot_receita*100:.1f}%)".replace(".", ","))

margem_fmt = pd.DataFrame(
    {
        "Mês": [meses_label[m] for m in meses],
        "Receita": margem_df["receita"].map(fmt_money),
        "Despesa": margem_df["despesa"].map(fmt_money),
        "Margem": margem_df["margem"].map(fmt_money),
        "Margem %": margem_df["margem_pct"].map(lambda v: f"{v:.1f}%".replace(".", ",")),
    }
)
st.dataframe(margem_fmt, hide_index=True, use_container_width=True)

with st.expander("Ver Margem sobre Receita GFP (efetivamente faturada)"):
    st.metric("Margem GFP", f"{fmt_money(tot_margem_gfp)} ({tot_margem_gfp/tot_receita_gfp*100:.1f}%)".replace(".", ","))
    margem_gfp_fmt = pd.DataFrame(
        {
            "Mês": [meses_label[m] for m in meses],
            "Receita GFP": margem_gfp_df["receita_gfp"].map(fmt_money),
            "Despesa": margem_gfp_df["despesa"].map(fmt_money),
            "Margem": margem_gfp_df["margem"].map(fmt_money),
            "Margem %": margem_gfp_df["margem_pct"].map(lambda v: f"{v:.1f}%".replace(".", ",")),
        }
    )
    st.dataframe(margem_gfp_fmt, hide_index=True, use_container_width=True)

# ---- Colaboradores ----
if f_colab:
    colabs, ativos = processar_colaboradores(f_colab.getvalue())
    st.subheader("Colaboradores")
    c1, c2, c3 = st.columns(3)
    c1.metric("Total ativos", len(ativos))
    c2.metric("Prodam", sum(1 for c in ativos if c.get("tipo") == "Prodam"))
    c3.metric("Terceiros", sum(1 for c in ativos if c.get("tipo") == "Terceiro"))

    por_nucleo = Counter(c.get("nucleo") for c in ativos)
    st.dataframe(
        pd.DataFrame(sorted(por_nucleo.items()), columns=["Núcleo", "Colaboradores"]),
        hide_index=True,
    )

    st.markdown("---")

# ---- Ordens de Serviço (dados_os_prodam) ----
if f_os:
    st.subheader("Ordens de Serviço Ativas")
    d_os = json.loads(f_os.getvalue())
    os_rows = []
    for chave in ("crp", "g4f"):
        info = d_os.get(chave)
        if info:
            os_rows.append(
                dict(
                    Fornecedor=info.get("fornecedor"),
                    **{
                        "N° OS": info.get("num_os"),
                        "SEI": info.get("sei"),
                        "TC": info.get("tc"),
                        "Gestor": info.get("gestor"),
                        "Valor UST": info.get("valor_ust"),
                    },
                )
            )
    if os_rows:
        st.dataframe(pd.DataFrame(os_rows), hide_index=True, use_container_width=True)
    else:
        st.caption("Nenhuma OS encontrada nas chaves esperadas (crp/g4f).")

    st.markdown("---")

# ---- Despesa terceiros (consumo CSV) ----
if f_consumo:
    st.subheader("Consumo / OS (bruto)")
    df_consumo = processar_consumo_csv(f_consumo.getvalue())
    st.dataframe(df_consumo, use_container_width=True)
    st.caption("Exibido conforme o CSV enviado — cálculo de despesa por fornecedor ainda não automatizado aqui.")

st.markdown("---")

# ---- Geração do Word (beta) ----
st.subheader("📄 Gerar relatório Word (beta)")
st.caption(
    "Preenche automaticamente, no template aprovado: KPIs, Resumo mensal GDS-1×GFP, Horas por "
    "núcleo, Saldo por contrato, Receita×Despesa×Margem (GDS-1 e GFP) e o Anexo 1 (detalhamento "
    "por contrato/mês). **Ainda não cobre**: Colaboradores e Fornecedores/OS no Word — essas "
    "seções eu ainda preencho junto com você aqui no chat."
)


def set_run_text(cell, para_idx, new_text, run_idx=0):
    p = cell.paragraphs[para_idx]
    runs = p.runs
    if not runs:
        p.add_run(new_text)
        return
    runs[run_idx].text = new_text
    for r in runs[run_idx + 1 :]:
        r.text = ""


def set_cell(table, row, col, text, para_idx=0):
    set_run_text(table.rows[row].cells[col], para_idx, text)


def fmt_pct(v):
    return f"{v:.1f}%".replace(".", ",")


def fmt_money_signed(v):
    return ("-" + fmt_money(-v)) if v < 0 else fmt_money(v)


def fmt_pct_signed(v):
    return ("-" + fmt_pct(-v)) if v < 0 else fmt_pct(v)


def encontrar_familia(texto_label: str, familias_calc: dict):
    """Casa o texto de um rótulo já existente no template (ex.: 'SMADS — TC85/SMADS/2023')
    com a família calculada correspondente, por normalização de secretaria+contrato."""
    alvo = normalizar(texto_label)
    for chave, f in familias_calc.items():
        if normalizar(f["nome"]) in alvo or alvo in normalizar(f["nome"]):
            return f
    return None


def gerar_docx(template_bytes, resumo, piv_nucleo, tot_posH, tot_posV, tot_gfpH, tot_gfpV,
               margem_df, margem_gfp_df, tot_despesa, familias_calc, meses):
    doc = Document(io.BytesIO(template_bytes))

    # Tabela 6: Resumo mensal
    try:
        t = doc.tables[6]
        for i, m in enumerate(resumo.index):
            row = i + 1
            r = resumo.loc[m]
            set_cell(t, row, 1, fmt_h(r["posH"]))
            set_cell(t, row, 2, fmt_money(r["posV"]))
            set_cell(t, row, 3, fmt_h(r["gfpH"]))
            set_cell(t, row, 4, fmt_money(r["gfpV"]))
        set_cell(t, 7, 1, fmt_h(tot_posH))
        set_cell(t, 7, 2, fmt_money(tot_posV))
        set_cell(t, 7, 3, fmt_h(tot_gfpH))
        set_cell(t, 7, 4, fmt_money(tot_gfpV))
    except Exception as e:
        st.warning(f"Tabela de Resumo mensal: {e}")

    # Tabela 7: por núcleo
    try:
        t = doc.tables[7]
        ordem = ["NSS1", "NSS2", "NSS3", "Cidadania"]
        tot_col = [0.0] * len(piv_nucleo.columns)
        for i, n in enumerate(ordem):
            if n not in piv_nucleo.index:
                continue
            row = i + 1
            vals = piv_nucleo.loc[n].values.tolist()
            for j, v in enumerate(vals[: len(piv_nucleo.columns)]):
                set_cell(t, row, j + 1, fmt_h(v))
                tot_col[j] += v
            set_cell(t, row, 7, fmt_h(sum(vals)))
        for j, v in enumerate(tot_col):
            set_cell(t, 5, j + 1, fmt_h(v))
        set_cell(t, 5, 7, fmt_h(sum(tot_col)))
    except Exception as e:
        st.warning(f"Tabela por núcleo: {e}")

    # KPIs
    tot_margem = tot_posV - tot_despesa
    tot_margem_gfp = tot_gfpV - tot_despesa
    for ti in (1, 14):
        try:
            t = doc.tables[ti]
            set_cell(t, 0, 0, fmt_money(tot_posV), para_idx=1)
            set_cell(t, 0, 2, fmt_pct_signed(tot_margem / tot_posV * 100), para_idx=1)
        except Exception:
            pass
    for ti in (2, 15):
        try:
            t = doc.tables[ti]
            set_cell(t, 0, 0, fmt_money(tot_gfpV), para_idx=1)
            set_cell(t, 0, 1, fmt_pct_signed(tot_margem_gfp / tot_gfpV * 100), para_idx=1)
        except Exception:
            pass

    # Tabela 16: Receita x Despesa x Margem mensal
    try:
        t = doc.tables[16]
        for i, m in enumerate(meses):
            row = i + 1
            r = margem_df.loc[m]
            set_cell(t, row, 1, fmt_money(r["receita"]))
            set_cell(t, row, 3, fmt_money_signed(r["margem"]))
            set_cell(t, row, 4, fmt_pct_signed(r["margem_pct"]))
        set_cell(t, 7, 1, fmt_money(tot_posV))
        set_cell(t, 7, 3, fmt_money_signed(tot_margem))
        set_cell(t, 7, 4, fmt_pct_signed(tot_margem / tot_posV * 100))
    except Exception as e:
        st.warning(f"Tabela Receita x Despesa x Margem: {e}")

    # Tabela 17: Margem sobre Receita GFP
    try:
        t = doc.tables[17]
        for i, m in enumerate(meses):
            row = i + 1
            r = margem_gfp_df.loc[m]
            set_cell(t, row, 1, fmt_money(r["receita_gfp"]))
            set_cell(t, row, 3, fmt_money_signed(r["margem"]))
            set_cell(t, row, 4, fmt_pct_signed(r["margem_pct"]))
        set_cell(t, 7, 1, fmt_money(tot_gfpV))
        set_cell(t, 7, 3, fmt_money_signed(tot_margem_gfp))
        set_cell(t, 7, 4, fmt_pct_signed(tot_margem_gfp / tot_gfpV * 100))
    except Exception as e:
        st.warning(f"Tabela Margem GFP: {e}")

    # Tabela 8: Saldo por contrato — casa cada linha pelo rótulo já existente na coluna 0
    try:
        t = doc.tables[8]
        tot_horasTotal = tot_acum = 0
        for row in t.rows[1:-1]:
            label = row.cells[0].text.strip()
            fam = encontrar_familia(label, familias_calc)
            if not fam:
                continue
            acum = sum(fam["por_mes"][m]["posH"] for m in meses)
            saldo = fam["horasTotal"] - acum
            pct = acum / fam["horasTotal"] * 100 if fam["horasTotal"] else 0
            set_run_text(row.cells[1], 0, fmt_h(round(fam["horasTotal"])))
            set_run_text(row.cells[2], 0, fmt_h(saldo))
            set_run_text(row.cells[3], 0, fmt_pct(pct))
            tot_horasTotal += fam["horasTotal"]
            tot_acum += acum
        ultima = t.rows[-1]
        set_run_text(ultima.cells[1], 0, fmt_h(round(tot_horasTotal)))
        set_run_text(ultima.cells[2], 0, fmt_h(tot_horasTotal - tot_acum))
        set_run_text(ultima.cells[3], 0, fmt_pct(tot_acum / tot_horasTotal * 100) if tot_horasTotal else "0,0%")
    except Exception as e:
        st.warning(f"Tabela Saldo por contrato: {e}")

    # Anexo 1: percorre o corpo do documento em ordem, casando cada parágrafo-título
    # (ex.: 'SMADS — TC85/SMADS/2023') com a tabela que vem logo depois dele.
    try:
        from docx.table import Table as _Table
        from docx.text.paragraph import Paragraph as _Paragraph

        body_items = []
        for child in doc.element.body.iterchildren():
            if child.tag.endswith("}p"):
                body_items.append(("p", _Paragraph(child, doc)))
            elif child.tag.endswith("}tbl"):
                body_items.append(("tbl", _Table(child, doc)))

        tot_geral = dict(gfpH=0.0, gfpV=0.0, posH=0.0, posV=0.0)
        pending_fam = None
        for kind, item in body_items:
            if kind == "p":
                fam = encontrar_familia(item.text.strip(), familias_calc)
                if fam:
                    pending_fam = fam
            elif kind == "tbl" and pending_fam is not None and len(item.columns) == 6:
                fam = pending_fam
                pending_fam = None
                acum = 0.0
                sub = dict(gfpH=0.0, gfpV=0.0, posH=0.0, posV=0.0)
                linhas_mes = item.rows[1 : 1 + len(meses)]
                if len(linhas_mes) != len(meses):
                    continue
                for i, m in enumerate(meses):
                    d = fam["por_mes"][m]
                    acum += d["posH"]
                    saldo = fam["horasTotal"] - acum
                    row = linhas_mes[i]
                    set_run_text(row.cells[1], 0, fmt_h(round(d["gfpH"], 2)))
                    set_run_text(row.cells[2], 0, fmt_money(d["gfpV"]))
                    set_run_text(row.cells[3], 0, fmt_h(round(d["posH"], 2)))
                    set_run_text(row.cells[4], 0, fmt_money(d["posV"]))
                    set_run_text(row.cells[5], 0, fmt_h(saldo))
                    for k in sub:
                        sub[k] += d[k]
                subtotal_row = item.rows[1 + len(meses)]
                set_run_text(subtotal_row.cells[1], 0, fmt_h(round(sub["gfpH"], 2)))
                set_run_text(subtotal_row.cells[2], 0, fmt_money(sub["gfpV"]))
                set_run_text(subtotal_row.cells[3], 0, fmt_h(round(sub["posH"], 2)))
                set_run_text(subtotal_row.cells[4], 0, fmt_money(sub["posV"]))
                set_run_text(subtotal_row.cells[5], 0, fmt_h(fam["horasTotal"] - acum))
                for k in tot_geral:
                    tot_geral[k] += sub[k]

        # TOTAL GERAL: última tabela de 5 colunas do documento (Anexo 1)
        for kind, item in reversed(body_items):
            if kind == "tbl" and len(item.columns) == 5 and len(item.rows) == 2:
                set_run_text(item.rows[1].cells[1], 0, fmt_h(round(tot_geral["gfpH"], 2)))
                set_run_text(item.rows[1].cells[2], 0, fmt_money(tot_geral["gfpV"]))
                set_run_text(item.rows[1].cells[3], 0, fmt_h(round(tot_geral["posH"], 2)))
                set_run_text(item.rows[1].cells[4], 0, fmt_money(tot_geral["posV"]))
                break
    except Exception as e:
        st.warning(f"Anexo 1 (detalhamento por contrato): {e}")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


with open("template_relatorio.docx", "rb") as f:
    template_bytes = f.read()

docx_bytes = gerar_docx(
    template_bytes,
    resumo,
    piv if not df_nucleo.empty else pd.DataFrame(),
    tot_posH, tot_posV, tot_gfpH, tot_gfpV,
    margem_df, margem_gfp_df, tot_despesa,
    familias_calc, meses,
)
st.download_button(
    "⬇️ Baixar relatório Word (parcialmente atualizado)",
    data=docx_bytes,
    file_name="Relatorio_Gestao_GDS1_atualizado.docx",
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
)

st.markdown("---")

# ---- Download dos dados calculados ----
st.subheader("Exportar dados")
buf = io.BytesIO()
with pd.ExcelWriter(buf, engine="openpyxl") as writer:
    resumo_fmt.to_excel(writer, sheet_name="Resumo mensal", index=False)
    if not df_nucleo.empty:
        piv_fmt.to_excel(writer, sheet_name="Por núcleo")
    pd.DataFrame(saldo_rows).to_excel(writer, sheet_name="Saldo por contrato", index=False)
    df.to_excel(writer, sheet_name="Detalhe bruto", index=False)
st.download_button(
    "⬇️ Baixar Excel com os dados corrigidos",
    data=buf.getvalue(),
    file_name="conferencia_gds1.xlsx",
    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
)
